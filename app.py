import io
import os
import re
import tempfile

import streamlit as st
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

from agents.graph import build_graph


def _md_to_docx(md_text: str) -> bytes:
    """Convert a Markdown report to a .docx file and return the bytes."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:])
            doc.add_paragraph(text, style="List Bullet")
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

load_dotenv()

SAMPLE_CSV = os.path.join(os.path.dirname(__file__), "sample_data", "churn.csv")
APP_ICON = os.path.join(os.path.dirname(__file__), "assets", "icon.png")
MAX_FILES = 5
MAX_FILE_SIZE_MB = 10
MAX_TOTAL_SIZE_MB = 25

# ── Validate API key on startup ──────────────────────────────────────────────

api_key = os.getenv("GOOGLE_API_KEY", "")
if not api_key:
    st.error(
        "**GOOGLE_API_KEY not found.** "
        "Add your Gemini key to the `.env` file in the project root.\n\n"
        "```\nGOOGLE_API_KEY=your_key_here\n```\n\n"
        "Get a free key at https://aistudio.google.com/apikey"
    )
    st.stop()

# ── Page config ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Multi-Agent CSV Analyst",
    page_icon=APP_ICON,
    layout="wide",
)

# ── Sidebar: session history ─────────────────────────────────────────────────

with st.sidebar:
    st.header("Session History")

    if "history" not in st.session_state:
        st.session_state.history = []

    if st.session_state.history:
        for item in reversed(st.session_state.history):
            with st.expander(f"Q: {item['question'][:50]}...", expanded=False):
                st.markdown(item["report"])
    else:
        st.caption("No queries yet. Ask a question to get started.")

# ── Main area ────────────────────────────────────────────────────────────────

st.title("Multi-Agent CSV Analyst")
st.caption(
    "Upload CSV files (up to 5 files, 10 MB each, 25 MB total), "
    "ask a question in plain English, and get a verified analysis report."
)

# ── Data source ──────────────────────────────────────────────────────────────

uploaded_files = st.file_uploader(
    "Upload CSV files",
    type=["csv"],
    accept_multiple_files=True,
    help=f"Up to {MAX_FILES} files, {MAX_FILE_SIZE_MB} MB each, {MAX_TOTAL_SIZE_MB} MB total.",
)

use_sample = st.button("Or load sample dataset (Telecom Churn)")

# Resolve CSV paths
csv_paths: list[str] = []

if uploaded_files:
    if len(uploaded_files) > MAX_FILES:
        st.error(f"Please upload at most {MAX_FILES} files. You uploaded {len(uploaded_files)}.")
        st.stop()

    total_bytes = sum(uf.size for uf in uploaded_files)
    oversized = [uf.name for uf in uploaded_files if uf.size > MAX_FILE_SIZE_MB * 1_048_576]

    if oversized:
        st.error(
            f"These files exceed the {MAX_FILE_SIZE_MB} MB limit: "
            f"{', '.join(oversized)}"
        )
        st.stop()

    if total_bytes > MAX_TOTAL_SIZE_MB * 1_048_576:
        st.error(
            f"Total upload size is {total_bytes / 1_048_576:.1f} MB. "
            f"Maximum allowed is {MAX_TOTAL_SIZE_MB} MB."
        )
        st.stop()

    for uf in uploaded_files:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv")
        tmp.write(uf.getvalue())
        tmp.close()
        csv_paths.append(tmp.name)

    names = ", ".join(f"**{uf.name}**" for uf in uploaded_files)
    total_mb = total_bytes / 1_048_576
    st.success(f"Loaded {len(uploaded_files)} file(s) ({total_mb:.1f} MB): {names}")

elif use_sample or st.session_state.get("using_sample"):
    csv_paths = [SAMPLE_CSV]
    st.session_state.using_sample = True
    st.success("Loaded **sample churn dataset** (100 rows, 16 columns)")

# ── Question input ───────────────────────────────────────────────────────────

question = st.text_area(
    "Ask a question about your data",
    placeholder="e.g. What is the churn rate by contract type?",
    height=80,
)

analyze_btn = st.button("Analyze", type="primary", use_container_width=True)

# ── Run the pipeline ─────────────────────────────────────────────────────────

if analyze_btn:
    if not csv_paths:
        st.error("Please upload CSV file(s) or load the sample dataset.")
        st.stop()
    if not question.strip():
        st.error("Please enter a question.")
        st.stop()

    graph = build_graph()

    initial_state = {
        "user_question": question.strip(),
        "df_path": csv_paths[0],
        "all_df_paths": csv_paths,
        "retry_count": 0,
        "generated_code": "",
        "code_output": "",
        "code_error": "",
        "report": "",
        "review_passed": False,
        "review_feedback": "",
        "dataset_info": "",
        "messages": [],
    }

    node_labels = {
        "profiler": "Profiling dataset...",
        "coder": "Generating Pandas code...",
        "executor": "Running code in sandbox...",
        "analyst": "Writing analysis report...",
        "reviewer": "Fact-checking the report...",
    }

    with st.status("Running multi-agent analysis...", expanded=True) as status:
        final_result = dict(initial_state)

        for step in graph.stream(initial_state, stream_mode="updates"):
            for node_name, node_output in step.items():
                label = node_labels.get(node_name, node_name)
                st.write(f"**{label}**")

                if node_name == "executor":
                    if node_output.get("code_error"):
                        st.warning(
                            f"Code failed (retry {node_output.get('retry_count', '?')}/3). "
                            "Sending error back to coder..."
                        )
                    else:
                        st.write("Code executed successfully.")

                if node_name == "reviewer":
                    if node_output.get("review_passed"):
                        st.write("Report approved.")
                    else:
                        st.warning(f"Review failed: {node_output.get('review_feedback', '')[:150]}")

                final_result.update(node_output)

        status.update(label="Analysis complete!", state="complete", expanded=False)

    # ── Display results ──────────────────────────────────────────────────

    st.divider()

    if final_result.get("report"):
        st.subheader("Analysis Report")
        st.markdown(final_result["report"])

        if not final_result.get("review_passed"):
            st.warning(
                "Note: The reviewer did not fully approve this report. "
                "Results may need manual verification."
            )

        docx_bytes = _md_to_docx(final_result["report"])

        st.download_button(
            label="Download Report (.docx)",
            data=docx_bytes,
            file_name="analysis_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.session_state.history.append({
            "question": question.strip(),
            "report": final_result["report"],
        })

    elif final_result.get("code_error"):
        st.error("The pipeline could not produce a result after all retries.")
        st.caption("Try rephrasing your question or checking your CSV format.")
    else:
        st.warning("No results were produced. Try rephrasing your question.")
